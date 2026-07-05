from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io, base64, tempfile
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from PIL import Image, ImageDraw, ImageFont

app = Flask(__name__)
CORS(app)

# ── Settings (set these in Render's Environment tab) ──────────────
GMAIL_ADDRESS = os.environ.get('smtp_USER', '')
GMAIL_APP_PASSWORD = os.environ.get('smtp_pass', '')
FIRM_EMAIL = os.environ.get('FIRM_EMAIL', 'info@tflaw.co.uk')
SOL_SIG_PATH = os.path.join(os.path.dirname(__file__), 'sol_sig.png')


def v(d, key):
    return str(d.get(key, '') or '').strip()


def shade_para(para, hex_col='1F3864'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    pPr.append(shd)


def section_hdr(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_para(p, '1F3864')
    run = p.add_run('  ' + text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def label_value_row(doc, label, value, label_w=6, value_w=10.5):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(label_w)
    table.columns[1].width = Cm(value_w)
    lc = table.rows[0].cells[0]
    vc = table.rows[0].cells[1]
    lc.width = Cm(label_w)
    vc.width = Cm(value_w)
    lc.paragraphs[0].add_run(label).font.size = Pt(9)
    vc.paragraphs[0].add_run(value).font.size = Pt(9)
    for cell in [lc, vc]:
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def decl_items(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        r = p.add_run('> ' + item)
        r.font.size = Pt(8.5)


def sig_block(doc, sig_label, sig_img_path, date_str, width_inches=3.0):
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    hl = table.rows[0].cells[0].paragraphs[0].add_run(sig_label)
    hl.bold = True
    hl.font.size = Pt(9)
    hr = table.rows[0].cells[1].paragraphs[0].add_run('Date:')
    hr.bold = True
    hr.font.size = Pt(9)
    sig_cell = table.rows[1].cells[0]
    sig_para = sig_cell.paragraphs[0]
    try:
        run = sig_para.add_run()
        run.add_picture(sig_img_path, width=Inches(width_inches))
    except Exception:
        sig_para.add_run('')
    table.rows[1].cells[1].paragraphs[0].add_run(date_str).font.size = Pt(9)
    table.rows[2].cells[0].paragraphs[0].add_run(' ')
    table.rows[2].cells[1].paragraphs[0].add_run(' ')
    doc.add_paragraph()


def make_client_sig_image(name, sig_data_url=None):
    """Use canvas signature if provided, otherwise render name as italic text."""
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    if sig_data_url and sig_data_url.startswith('data:image'):
        header, b64data = sig_data_url.split(',', 1)
        img_bytes = base64.b64decode(b64data)
        with open(tmp.name, 'wb') as f:
            f.write(img_bytes)
    else:
        img = Image.new('RGB', (500, 90), color='white')
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf', 32)
        except Exception:
            font = ImageFont.load_default()
        draw.text((10, 25), name, fill=(13, 27, 42), font=font)
        img.save(tmp.name, format='PNG')
    return tmp.name


def make_transfer(d, client_sig_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = 'Table Grid'
    hdr_table.rows[0].cells[0].paragraphs[0].add_run(
        'Civil Legal Aid\nLegal Aid Online Transfer Declaration').font.size = Pt(11)
    hdr_table.rows[0].cells[1].paragraphs[0].add_run('April 2018\nCIVTR/LAO').font.size = Pt(8)

    doc.add_paragraph()
    section_hdr(doc, 'A.  Key Details')
    label_value_row(doc, "Solicitor's name:", v(d, 'solname'))
    label_value_row(doc, 'Legal aid reference number:', v(d, 'laref'))
    label_value_row(doc, "Applicant's full name:", v(d, 'appname'))
    label_value_row(doc, 'Reason for transfer:', v(d, 'reason'))

    doc.add_paragraph()
    section_hdr(doc, "B.  Applicant or Representative's Declaration")
    decl_items(doc, [
        'I want the solicitor named at Section A to act for me in this case.',
        'I ask you to transfer my legal aid certificate to that solicitor.',
        'I give my permission to SLAB to notify the details of this transfer request to the solicitor named in my legal aid certificate before any transfer, to allow SLAB to check the circumstances, before deciding whether to transfer legal aid.',
        'I confirm the information given by me in this application is correct.',
        'I confirm that I have read and agree that the reasons provided are my reasons for seeking a transfer of legal aid.',
        'I understand that the information I have provided on this form may be used by SLAB or other bodies responsible for auditing or administering public funds for the prevention and detection of fraud or abuse of legal aid.',
    ])
    sig_block(doc, 'Signature of applicant/representative:', client_sig_path, v(d, 'sigdate'), width_inches=3.0)

    section_hdr(doc, "C.  Solicitor's Declaration")
    decl_items(doc, [
        'Whether subject to the transfer of legal aid or not, I have agreed to act for the Applicant.',
        'Any opinion expressed by me in the application represents my professional opinion.',
        'I certify that to the best of my knowledge and belief, the information given is correct. I consent to the disclosure of the application, associated documentation and client case file for quality assurance, including peer review and stage reporting purposes, at any stage during or after the proceedings.',
    ])
    sig_block(doc, 'Signature of solicitor:', SOL_SIG_PATH, v(d, 'sigdate'), width_inches=3.0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def send_transfer_email(d, transfer_buf):
    appname = v(d, 'appname')
    ref = v(d, 'ref')
    body = ('New Civil Legal Aid Transfer Declaration submitted online.\n\n'
            'Applicant: ' + appname + '\n'
            'Incoming Solicitor: ' + v(d, 'solname') + '\n'
            'Legal Aid Reference: ' + v(d, 'laref') + '\n'
            'Reason for Transfer: ' + v(d, 'reason') + '\n'
            'Date Signed: ' + v(d, 'sigdate') + '\n'
            'Reference: ' + ref + '\n\n'
            'The completed transfer declaration is attached as a Word document.\n'
            'Please countersign and submit to SLAB via Legal Aid Online.')

    msg = MIMEMultipart()
    msg['From'] = GMAIL_ADDRESS
    msg['To'] = FIRM_EMAIL
    msg['Subject'] = 'New Legal Aid Transfer Declaration - ' + appname + ' - ' + ref
    msg.attach(MIMEText(body, 'plain'))

    part = MIMEBase('application', 'octet-stream')
    part.set_payload(transfer_buf.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment; filename=CIVTR_LAO_' + appname.replace(' ', '_') + '.docx')
    msg.attach(part)

    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
    server.send_message(msg)
    server.quit()


@app.route('/submit', methods=['POST'])
def submit_transfer():
    try:
        d = request.json
        d['ref'] = 'TFL-' + str(int(datetime.now().timestamp()))[-6:]
        print('Transfer Submit:', v(d, 'appname'))

        client_sig_path = make_client_sig_image(v(d, 'appname'), d.get('sigDataURL', ''))

        transfer_buf = make_transfer(d, client_sig_path)
        print('Transfer doc built')

        send_transfer_email(d, transfer_buf)
        print('Transfer email sent')

        os.unlink(client_sig_path)
        return jsonify({'ok': True, 'ref': d['ref']})
    except Exception as e:
        import traceback
        print('ERROR:', str(e))
        traceback.print_exc()
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/')
def home():
    return 'Transfer Declaration Server - OK'


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
